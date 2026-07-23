from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Maestro_Codebase_Architecture_Report.docx"
BLUE, NAVY, INK, MUTED, LIGHT, PALE, GREEN = "2E74B5", "1F4D78", "17212B", "5F6B76", "F2F4F7", "E8EEF5", "247A4A"


def shade(cell, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell_width(cell, width):
    tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side, value in (("top", 80), ("start", 120), ("bottom", 80), ("end", 120)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header(row):
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    row._tr.get_or_add_trPr().append(header)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, text in enumerate(headers):
        shade(table.rows[0].cells[index], LIGHT)
        run = table.rows[0].cells[index].paragraphs[0].add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(NAVY)
    repeat_header(table.rows[0])
    for values in rows:
        cells = table.add_row().cells
        for index, text in enumerate(values):
            cells[index].paragraphs[0].add_run(text)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    shade(table.cell(0, 0), PALE)
    paragraph = table.cell(0, 0).paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    label_run = paragraph.add_run(f"{label.upper()}  ")
    label_run.bold = True
    label_run.font.color.rgb = RGBColor.from_string(BLUE)
    paragraph.add_run(text)
    repeat_header(table.rows[0])
    set_table_geometry(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def set_page_field(paragraph):
    paragraph.add_run("Page ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


doc = Document()
section = doc.sections[0]
section.page_width, section.page_height = Inches(8.5), Inches(11)
section.top_margin = section.right_margin = section.bottom_margin = section.left_margin = Inches(1)
section.header_distance = section.footer_distance = Inches(0.492)

normal = doc.styles["Normal"]
normal.font.name, normal.font.size = "Calibri", Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after, normal.paragraph_format.line_spacing = Pt(6), 1.10
for name, size, before, after, color in (
    ("Title", 26, 0, 8, INK), ("Subtitle", 13, 0, 16, MUTED),
    ("Heading 1", 16, 16, 8, BLUE), ("Heading 2", 13, 12, 6, BLUE),
    ("Heading 3", 12, 8, 4, NAVY),
):
    style = doc.styles[name]
    style.font.name, style.font.size = "Calibri", Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = name != "Subtitle"
    style.paragraph_format.space_before, style.paragraph_format.space_after = Pt(before), Pt(after)
list_style = doc.styles["List Bullet"]
list_style.font.name, list_style.font.size = "Calibri", Pt(11)
list_style.paragraph_format.space_after, list_style.paragraph_format.line_spacing = Pt(8), 1.167

header = section.header.paragraphs[0]
header.text = "MAESTRO  /  CODEBASE ARCHITECTURE"
header.runs[0].font.size, header.runs[0].bold = Pt(9), True
header.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_page_field(footer)
for run in footer.runs:
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MUTED)

kicker = doc.add_paragraph()
kicker.paragraph_format.space_before = Pt(18)
run = kicker.add_run("ARCHITECTURE & RELEASE READINESS BRIEF")
run.bold, run.font.size = True, Pt(10)
run.font.color.rgb = RGBColor.from_string(BLUE)
doc.add_paragraph("Maestro Codebase Architecture", style="Title")
doc.add_paragraph("Clean package boundaries, reusable UI composition, validation evidence, and staged evolution", style="Subtitle")
add_table(doc, ["Repository", "Branch", "Prepared"], [["subhangR/agent-maestro", "staging", date.today().strftime("%d %B %Y")]], [3900, 1860, 3600])
add_callout(doc, "Outcome", "The monorepo now has an explicit architecture manifest, enforceable cross-package dependency rules, documented frontend/backend/shared groupings, and a reusable session-view feature split into state and presentation modules. Stable package paths were preserved to avoid deployment and release regressions.")

doc.add_heading("1. Executive decision", level=1)
doc.add_paragraph("Keep the current package-based monorepo structure and treat logical grouping as the public architecture contract. Do not physically rename mature top-level packages during an active staging cycle. The repository now expresses those groups in architecture.json and validates them with scripts/check-architecture.mjs.")
add_bullets(doc, [
    "Preserve compatibility for Tauri configuration, build scripts, deployment automation, documentation links, and downstream tooling.",
    "Enforce dependency direction continuously instead of relying on directory appearance alone.",
    "Extract reusable frontend behavior from large orchestration components as features change.",
    "Move to apps/ and packages/ only through a dedicated migration with compatibility shims and full release validation.",
])

doc.add_heading("2. Logical package map", level=1)
add_table(doc, ["Group", "Packages", "Owned responsibility"], [
    ["Frontend", "maestro-ui; maestro-web; maestro-mobile; website", "Desktop/browser/mobile presentation, interaction, and platform UI."],
    ["Backend", "maestro-server; maestro-gateway; functions", "APIs, orchestration, persistence, realtime transport, auth, and cloud adapters."],
    ["Interfaces", "maestro-cli", "Command surface for users, agents, and automation."],
    ["Shared", "maestro-pty-protocol", "Runtime-independent PTY contracts and parsing."],
    ["Operations", "deploy; scripts", "Build, deployment, setup, diagnostics, and maintenance."],
], [1440, 3600, 4320])

doc.add_heading("3. Dependency boundaries", level=1)
doc.add_paragraph("The dependency rule is intentionally simple: applications depend on shared contracts; clients and servers communicate across runtime boundaries rather than importing one another’s implementations.")
add_table(doc, ["Source", "Allowed", "Prohibited"], [
    ["Frontend", "Shared contracts, UI libraries, platform adapters", "Backend implementation packages"],
    ["Backend", "Domain/application contracts, infrastructure adapters", "Frontend implementation packages"],
    ["Shared", "Dependency-free protocol/types where practical", "Application or runtime packages"],
    ["Server domain", "Domain values, repository/event interfaces", "Express, filesystem, WebSocket, cloud SDKs"],
], [1800, 3780, 3780])
add_callout(doc, "Automated gate", "bun run check:architecture verifies every declared path, prevents duplicate grouping, resolves workspace package names, and rejects frontend↔backend or shared→application package dependencies.")

doc.add_heading("4. Frontend decomposition implemented", level=1)
doc.add_paragraph("The pending session workspace improvement previously placed responsive defaults, local-storage persistence, view-state derivation, and three-button rendering directly inside AppWorkspace. That behavior is now decomposed into focused units:")
add_table(doc, ["Module", "Responsibility"], [
    ["useSessionViewMode.ts", "Owns the chat/split/terminal state, responsive initial value, preference validation, persistence, and derived visibility flags."],
    ["SessionViewSelector.tsx", "Pure reusable presentation for the accessible tab selector."],
    ["AppWorkspace.tsx", "Composes workspace data, terminals, activity, and the extracted feature without owning its internal policy."],
    ["useSessionViewMode.test.ts", "Covers narrow-screen defaults, valid preference restoration, and invalid preference fallback."],
], [3000, 6360])

doc.add_heading("5. Backend decomposition retained", level=1)
doc.add_paragraph("maestro-server already follows a credible inward dependency model. API routes adapt HTTP inputs; application services own use cases; domain interfaces define repositories/events; infrastructure implements filesystem, auth, WebSocket, skills, and external integrations; the container composes the graph. This structure should be extended rather than replaced.")
add_bullets(doc, [
    "Keep validation and transport mapping in API adapters.",
    "Keep use-case decisions and orchestration in application services.",
    "Keep business contracts free of Express, storage, and cloud SDK concerns.",
    "Add new persistence or realtime implementations behind existing interfaces.",
])

doc.add_heading("6. Verification evidence", level=1)
add_table(doc, ["Check", "Result", "Notes"], [
    ["Architecture boundary check", "PASS", "11 grouped paths and 8 package manifests validated."],
    ["Frontend test suite", "PASS", "81 test files; 656 tests passed."],
    ["Frontend TypeScript + Vite build", "PASS", "7,226 modules transformed; production bundle completed with 4 GB Node heap."],
    ["Server TypeScript build", "PASS", "tsc completed."],
    ["CLI TypeScript build", "PASS", "tsc completed."],
    ["Diff whitespace check", "PASS", "No whitespace errors reported."],
], [3600, 1320, 4440])
doc.add_paragraph("Non-blocking observations: the UI bundle reports large chunks and static/dynamic import overlap; these are optimization opportunities rather than correctness failures. The root build now supplies the heap needed by the current bundle graph.")

doc.add_heading("7. Safe migration roadmap", level=1)
add_table(doc, ["Stage", "Action", "Exit criterion"], [
    ["Now", "Use architecture.json as the canonical grouping; require check:architecture in local/CI validation.", "No cross-layer package dependency violations."],
    ["Next", "Continue feature extraction from large UI components into feature folders, hooks, services, and platform adapters.", "New behavior has focused tests and narrow public APIs."],
    ["Then", "Extract additional shared API/event schemas only where at least two runtimes consume them.", "Contracts remain runtime-independent and versioned."],
    ["Dedicated migration", "Optionally move physical paths to apps/ and packages/ with scripted reference updates and compatibility shims.", "Desktop, web, CLI, server, gateway, mobile, CI, and deploy smoke tests pass."],
], [1680, 4800, 2880])

doc.add_heading("8. Release checklist", level=1)
add_bullets(doc, [
    "Run bun run check:architecture.",
    "Run the primary frontend, server, CLI, and shared-protocol tests.",
    "Run bun run build:all and review Vite warnings for new regressions.",
    "Confirm git diff --check and review all staged files before commit.",
    "Push staging only after GitHub authentication succeeds and verify the staging deployment workflow.",
])
add_callout(doc, "Current publishing constraint", "The local staging branch can be committed and is ready to push, but both configured GitHub CLI credentials were invalid at audit time. Re-authentication with gh auth login is required before the remote staging branch can be updated.")

doc.core_properties.title = "Maestro Codebase Architecture Report"
doc.core_properties.subject = "Frontend/backend decomposition and release readiness"
doc.core_properties.author = "Maestro Engineering"
doc.core_properties.keywords = "Maestro, architecture, frontend, backend, monorepo, staging"
doc.save(OUT)
print(OUT)
