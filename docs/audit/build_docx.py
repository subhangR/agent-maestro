from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont
import re

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Agent_Maestro_Product_Functional_Architecture_Audit.docx"
NAVY = "13233A"
BLUE = "2667FF"
CYAN = "19B5C5"
INK = "18212F"
MUTED = "5E6B7A"
PALE = "EEF4FF"
LINE = "D7E0EA"

def font(size, bold=False):
    candidates = ["/System/Library/Fonts/Supplemental/Arial.ttf", "/Library/Fonts/Arial.ttf"]
    for p in candidates:
        if Path(p).exists():
            return ImageFont.truetype(p, size=size)
    return ImageFont.load_default()

def rounded(draw, xy, fill, outline=None, radius=24, width=3):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)

def make_architecture(path):
    im = Image.new("RGB", (1800, 1000), "#F7F9FC")
    d = ImageDraw.Draw(im)
    d.text((70, 42), "Agent Maestro — system landscape", fill="#13233A", font=font(54, True))
    cols = [(70, "EXPERIENCE", "#E7EEFF"), (510, "ORCHESTRATION", "#E8F7F8"), (950, "PERSISTENCE", "#F2ECFF"), (1390, "OPERATIONS", "#FFF1E5")]
    boxes = [
        ["Tauri desktop", "Browser UI", "Mobile", "CLI"],
        ["Express REST", "Event WebSocket", "PTY host", "Agent providers"],
        ["Local JSON", "Session logs", "Firestore", "RTDB + FCM"],
        ["Gateway", "nginx / TLS", "systemd", "Firebase Hosting"]]
    for ci, (x, title, bg) in enumerate(cols):
        rounded(d, (x, 130, x+360, 870), bg, "#D0D9E6", 30, 3)
        d.text((x+28, 165), title, fill="#435269", font=font(25, True))
        for i, label in enumerate(boxes[ci]):
            y = 245 + i*135
            rounded(d, (x+28, y, x+332, y+88), "#FFFFFF", "#BCC8D8", 18, 2)
            d.text((x+52, y+27), label, fill="#18212F", font=font(25))
    for x1, x2 in [(430,510),(870,950),(1310,1390)]:
        d.line((x1,500,x2,500), fill="#2667FF", width=8)
        d.polygon([(x2,500),(x2-20,486),(x2-20,514)], fill="#2667FF")
    d.text((70, 918), "Core orchestration remains local-first; Firebase is the optional collaboration plane.", fill="#5E6B7A", font=font(27))
    im.save(path, quality=95)

def make_flow(path):
    im = Image.new("RGB", (1800, 760), "white")
    d = ImageDraw.Draw(im)
    d.text((70, 45), "From task intent to verified agent work", fill="#13233A", font=font(54, True))
    labels = ["Choose project\nand task", "Review agent, model,\ncwd and permissions", "Create session\nand manifest", "Open PTY and\nlaunch provider", "Stream events, logs\nand progress", "Verify outcome\nand complete task"]
    for i, label in enumerate(labels):
        x = 55 + i*290
        rounded(d, (x, 205, x+245, 455), "#EEF4FF" if i%2==0 else "#E8F7F8", "#B8C8E5", 28, 3)
        d.ellipse((x+88, 235, x+158, 305), fill="#2667FF")
        n = str(i+1)
        d.text((x+110, 246), n, fill="white", font=font(32, True))
        for j, line in enumerate(label.split("\n")):
            d.text((x+22, 335+j*38), line, fill="#18212F", font=font(22))
        if i < len(labels)-1:
            d.line((x+245,330,x+290,330), fill="#2667FF", width=7)
            d.polygon([(x+290,330),(x+273,318),(x+273,342)], fill="#2667FF")
    d.text((70, 575), "Human control points", fill="#2667FF", font=font(25, True))
    d.text((70, 615), "Scope • permissions • destructive actions • completion evidence • recovery", fill="#435269", font=font(28))
    im.save(path, quality=95)

def make_roadmap(path):
    im = Image.new("RGB", (1800, 720), "#F7F9FC")
    d = ImageDraw.Draw(im)
    d.text((70, 42), "Roadmap: coherence before expansion", fill="#13233A", font=font(54, True))
    phases = [
        ("0–30 days", "STABILIZE & EXPLAIN", ["Reconcile docs", "Threat model", "Core-flow smoke tests"]),
        ("31–90 days", "SIMPLIFY & HARDEN", ["State ownership", "Accessible design system", "SQLite repository"]),
        ("3–6 months", "OPERATE CONFIDENTLY", ["PTY isolation", "Observability + SLOs", "Signed releases"]),
    ]
    colors = ["#E7EEFF", "#E8F7F8", "#F2ECFF"]
    for i, (time, title, items) in enumerate(phases):
        x = 70 + i*570
        rounded(d, (x,150,x+510,610), colors[i], "#C1CDDD", 30, 3)
        d.text((x+32,185), time, fill="#2667FF", font=font(28, True))
        d.text((x+32,240), title, fill="#13233A", font=font(26, True))
        for j,item in enumerate(items):
            y=335+j*70
            d.ellipse((x+35,y+7,x+53,y+25),fill="#19B5C5")
            d.text((x+72,y),item,fill="#18212F",font=font(24))
    im.save(path, quality=95)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=tcMar.find(qn('w:'+m))
        if node is None: node=OxmlElement('w:'+m); tcMar.append(node)
        node.set(qn('w:w'),str(v)); node.set(qn('w:type'),'dxa')

def add_page_number(paragraph):
    run=paragraph.add_run(); fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); run._r.addnext(fld)

def set_run(run, size=None, color=INK, bold=None, italic=None):
    run.font.name='Arial'; run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'),'Arial'); run._element.rPr.rFonts.set(qn('w:hAnsi'),'Arial')
    if size: run.font.size=Pt(size)
    run.font.color.rgb=RGBColor.from_string(color)
    if bold is not None: run.bold=bold
    if italic is not None: run.italic=italic

def add_rich_paragraph(doc, text, style=None):
    p=doc.add_paragraph(style=style)
    for part in re.split(r'(\*\*.*?\*\*|`.*?`|https?://\S+)', text):
        if not part: continue
        if part.startswith('**') and part.endswith('**'):
            r=p.add_run(part[2:-2]); set_run(r,bold=True)
        elif part.startswith('`') and part.endswith('`'):
            r=p.add_run(part[1:-1]); set_run(r,size=9,color=NAVY); r.font.name='Menlo'
        else:
            r=p.add_run(part); set_run(r)
    return p

def parse_md(doc, path, images):
    lines=path.read_text().splitlines(); in_code=False; code=[]
    for line in lines:
        if line.startswith('```'):
            if in_code:
                title='Diagram specification' if code and code[0] in ('flowchart TB','flowchart LR','sequenceDiagram') else 'Technical example'
                p=doc.add_paragraph(); r=p.add_run(title); set_run(r,9,MUTED,True)
                for c in code[:12]:
                    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.25); p.paragraph_format.space_after=Pt(1)
                    r=p.add_run(c); set_run(r,8,NAVY); r.font.name='Menlo'
                code=[]; in_code=False
            else: in_code=True
            continue
        if in_code: code.append(line); continue
        if not line.strip(): continue
        if line.startswith('# '):
            doc.add_page_break(); p=doc.add_paragraph(line[2:],style='Heading 1')
            if 'Functional' in line and 'flow' in images: doc.add_picture(str(images['flow']),width=Inches(6.45))
            if 'System architecture' in line and 'architecture' in images: doc.add_picture(str(images['architecture']),width=Inches(6.45))
            if 'roadmap' in line.lower() and 'roadmap' in images: doc.add_picture(str(images['roadmap']),width=Inches(6.45))
        elif line.startswith('## '): doc.add_paragraph(line[3:],style='Heading 2')
        elif line.startswith('### '): doc.add_paragraph(line[4:],style='Heading 3')
        elif re.match(r'^\d+\. ',line): add_rich_paragraph(doc,re.sub(r'^\d+\. ','',line),style='List Number')
        elif line.startswith('- '): add_rich_paragraph(doc,line[2:],style='List Bullet')
        else: add_rich_paragraph(doc,line)

def build():
    arch=ROOT/'architecture.png'; flow=ROOT/'session_flow.png'; roadmap=ROOT/'roadmap.png'
    make_architecture(arch); make_flow(flow); make_roadmap(roadmap)
    doc=Document(); sec=doc.sections[0]
    sec.top_margin=Inches(.8); sec.bottom_margin=Inches(.75); sec.left_margin=Inches(1); sec.right_margin=Inches(1)
    sec.header_distance=Inches(.35); sec.footer_distance=Inches(.35)
    styles=doc.styles
    normal=styles['Normal']; normal.font.name='Arial'; normal.font.size=Pt(10.5); normal.font.color.rgb=RGBColor.from_string(INK)
    normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.15
    for name,size,color,before,after in [('Title',30,NAVY,0,10),('Subtitle',14,MUTED,0,10),('Heading 1',20,NAVY,12,8),('Heading 2',15,BLUE,14,6),('Heading 3',11,CYAN,10,4)]:
        st=styles[name]; st.font.name='Arial'; st.font.size=Pt(size); st.font.bold=name!='Subtitle'; st.font.color.rgb=RGBColor.from_string(color)
        st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True
    for n in ['List Bullet','List Number']:
        styles[n].font.name='Arial'; styles[n].font.size=Pt(10.5); styles[n].paragraph_format.space_after=Pt(4)
    hp=sec.header.paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    r=hp.add_run('AGENT MAESTRO  /  PRODUCT • FUNCTION • ARCHITECTURE'); set_run(r,8,MUTED,True)
    fp=sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    r=fp.add_run('AUDITED STAGING EDITION  •  '); set_run(r,8,MUTED,True); add_page_number(fp)

    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(95); r=p.add_run('AGENT MAESTRO'); set_run(r,11,BLUE,True)
    p=doc.add_paragraph(style='Title'); p.add_run('Product, Functional &\nArchitecture Audit')
    p=doc.add_paragraph(style='Subtitle'); p.add_run('A plain-English reference for product leaders, designers, engineers, operators, and new contributors')
    doc.add_picture(str(arch),width=Inches(6.45))
    t=doc.add_table(rows=3,cols=2); t.autofit=False
    facts=[('Reviewed branch','staging'),('Audit date','24 July 2026'),('Scope','2,516 tracked paths • ~647k lines including assets/docs')]
    for i,(a,b) in enumerate(facts):
        t.columns[0].width=Inches(1.45); t.columns[1].width=Inches(5.0)
        for c in t.rows[i].cells: set_cell_margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade(t.cell(i,0),PALE); t.cell(i,0).text=a; t.cell(i,1).text=b
        for rr in t.cell(i,0).paragraphs[0].runs: set_run(rr,9,BLUE,True)
        for rr in t.cell(i,1).paragraphs[0].runs: set_run(rr,9,INK)
    doc.add_page_break(); doc.add_paragraph('How to use this report',style='Heading 1')
    add_rich_paragraph(doc,'This report separates verified current-state behavior from recommendations. The code remains the source of truth; older repository documentation is treated as evidence only when it matches implementation.')
    for x in ['Product & UX — what Maestro is and how it should feel.','Functional specification — what each capability must do.','System architecture — how clients, server, storage, Firebase, PTY, gateway and hosting interact.','Audit & roadmap — what is strong, what is risky, and what to do next.']:
        add_rich_paragraph(doc,x,style='List Bullet')
    doc.add_paragraph('Executive takeaway',style='Heading 2')
    p=doc.add_paragraph(); shade_dummy=None
    r=p.add_run('Maestro should compete on trustworthy control and legibility of multi-agent work—not on adding the largest number of features.'); set_run(r,13,NAVY,True)
    files=['01_PRODUCT_AND_UX.md','02_FUNCTIONAL_SPECIFICATION.md','03_SYSTEM_ARCHITECTURE.md','04_CODEBASE_AUDIT_AND_ROADMAP.md']
    images={'architecture':arch,'flow':flow,'roadmap':roadmap}
    for f in files: parse_md(doc,ROOT/f,images)
    doc.add_page_break(); doc.add_paragraph('Source and evidence register',style='Heading 1')
    sources=[
        ('Repository','https://github.com/subhangR/agent-maestro'),
        ('Apple HIG','https://developer.apple.com/design/human-interface-guidelines/'),
        ('Apple accessibility','https://developer.apple.com/design/human-interface-guidelines/accessibility'),
        ('Material Design 3','https://m3.material.io/'),
        ('Meta design resources','https://design.facebook.com/toolsandresources/'),
        ('Uber Base','https://base.uber.com/'),
        ('Firebase documentation','https://firebase.google.com/docs'),
        ('Tauri documentation','https://v2.tauri.app/'),
    ]
    for a,b in sources: add_rich_paragraph(doc,f'**{a}:** {b}',style='List Bullet')
    doc.add_paragraph('Audit note',style='Heading 2')
    add_rich_paragraph(doc,'The review inventoried every tracked path and semantically traced the runtime-bearing code. Binary images, audio, fonts, lockfiles, generated output, and historical/reference documents were cataloged rather than falsely described as line-by-line business logic.')
    doc.save(OUT)
    print(OUT)

if __name__=='__main__': build()
